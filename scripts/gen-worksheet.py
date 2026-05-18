"""
Build worksheet.docx from data/content.json.
Re-runnable: regenerates the file from the current content so the worksheet
always matches the website.
"""
import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT = Path("/Users/hongyu/year7-space-explorer")
CONTENT = json.loads((PROJECT / "data/content.json").read_text())["celestialBodies"]
OUT = PROJECT / "worksheet.docx"

doc = Document()

# Page setup — A4, 1.5 cm margins for max writing space
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

# Default font: Calibri 11pt
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

NAVY = RGBColor(0x1a, 0x3a, 0x6a)
MID  = RGBColor(0x2a, 0x4a, 0x7a)
GREY = RGBColor(0x55, 0x55, 0x55)


def add_heading(text, level=1, color=NAVY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(20)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(8 if level > 1 else 0)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_subtitle(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = GREY
    run.font.size = Pt(10.5)
    run.italic = True
    p.paragraph_format.space_after = Pt(8)
    return p


def add_field_line(label, width_cm=4):
    """A 'Label: ____________' field for filling in."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    return p


# ── Title block ──
add_heading("Celestial Bodies", level=1)
add_subtitle("Note-taking worksheet — Year 7 Science")

# Name / Class / Date row
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
for label in ["Name:  ", "Class:  ", "Date:  "]:
    run = p.add_run(label)
    run.bold = True
    p.add_run("_" * 18 + "   ")

# Horizontal rule (paragraph with bottom border)
def add_rule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '2A4A7A')
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)
add_rule()


# ── Learning intentions ──
add_heading("Learning Intentions", level=2)
for li in CONTENT["learningIntentions"]:
    p = doc.add_paragraph(li, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)

# ── Instructions ──
add_heading("Instructions", level=2)
inst = doc.add_paragraph()
inst.add_run("Open the website's ")
r = inst.add_run("Celestial Bodies"); r.bold = True
inst.add_run(" tab and read the description for each object. Use the ")
r = inst.add_run("word bank"); r.bold = True
inst.add_run(
    " for each section to fill in the blanks. Then answer the comparison questions in your own words and complete the self-check at the end."
)


def add_word_bank_paragraph(terms):
    """A pill-style word bank: [term] [term] [term]"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    label = p.add_run("Word bank:  ")
    label.bold = True
    label.font.color.rgb = NAVY
    for i, t in enumerate(terms):
        if i > 0:
            p.add_run("    ")
        run = p.add_run(f"  {t}  ")
        run.font.color.rgb = MID
        # Outline the term by using a light shading box (via underline)
        run.underline = True
    return p


def make_blanked_text(description, blanks):
    """Replace each key term with an underlined-blank of matching length."""
    out = description
    # Each term replaced ONCE — same as the HTML page
    for term in blanks:
        pattern = r'\b(' + re.escape(term) + r')\b'
        # Use a marker we can split on later (avoid double replacement)
        out, n = re.subn(pattern, '⟦BLANK⟧', out, count=1)
    return out


def add_blanked_paragraph(text):
    """text contains ⟦BLANK⟧ markers; turn each into an underlined blank line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.6  # extra room so blanks look fillable
    parts = text.split('⟦BLANK⟧')
    for i, chunk in enumerate(parts):
        if chunk:
            p.add_run(chunk)
        if i < len(parts) - 1:
            # The blank itself — underscored space wide enough to write in
            blank_run = p.add_run('  ' + '_' * 14 + '  ')
            blank_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p


def add_notes_lines(count=2):
    """Empty underlined lines for free-form notes."""
    label = doc.add_paragraph()
    r = label.add_run("My own notes:")
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = GREY
    label.paragraph_format.space_after = Pt(2)
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:color'), 'AAAAAA')
        bottom.set(qn('w:space'), '1')
        pBdr.append(bottom)
        pPr.append(pBdr)
        # Empty content; the border supplies the line
        p.add_run(" ")


# ── Body sections ──
add_heading("The eight celestial bodies", level=2)
add_subtitle("Fill in the blanks in each description, then jot any extra notes you want.")

for item in CONTENT["items"]:
    # Card-style heading row
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(item.get("icon", "•") + "  ")
    r.font.size = Pt(16)
    r = p.add_run(item["name"])
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = NAVY

    # Word bank
    add_word_bank_paragraph(sorted(set(item["wordBank"])))

    # Blanked description
    blanked = make_blanked_text(item["description"], item["blanks"])
    add_blanked_paragraph(blanked)

    # Examples
    ex = doc.add_paragraph()
    ex.paragraph_format.space_after = Pt(2)
    r = ex.add_run("Examples: ")
    r.bold = True
    r.font.color.rgb = NAVY
    r.font.size = Pt(10)
    r = ex.add_run(item["examples"])
    r.font.size = Pt(10)
    r.font.color.rgb = GREY

    # Notes lines
    add_notes_lines(2)


# ── Comparison questions ──
doc.add_page_break()
add_heading("Spot the difference", level=2)
add_subtitle("Use what you've learned to answer each question in your own words.")
for i, c in enumerate(CONTENT["comparisons"], start=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{i}. ")
    r.bold = True
    r.font.color.rgb = NAVY
    r = p.add_run(c["q"])
    r.bold = True
    # 3 blank lines
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:color'), 'AAAAAA')
        bottom.set(qn('w:space'), '1')
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.add_run(" ")


# ── Self-check ──
add_heading("Self-check", level=2)
add_subtitle("Tick a box when you can answer it confidently in your own words.")
for item in CONTENT["items"]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("☐  ")
    r.font.size = Pt(13)
    p.add_run("I can describe in my own words what a ")
    r = p.add_run(item["name"].lower().rstrip("s") if item["name"].endswith("s") else item["name"].lower())
    r.bold = True
    p.add_run(" is.")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("☐  ")
r.font.size = Pt(13)
p.add_run("I can explain the difference between any two of these eight bodies.")


# ── Answer key (separate page) ──
doc.add_page_break()
add_heading("Teacher answer key", level=2)
add_subtitle("The words that belong in each blank, in order. Detach this page before handing out.")
for item in CONTENT["items"]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{item['name']}: ")
    r.bold = True
    r.font.color.rgb = NAVY
    r = p.add_run(", ".join(item["blanks"]))
    r.font.color.rgb = RGBColor(0xb2, 0x10, 0x30)

# Comparison sample answers
add_heading("Sample answers — Spot the difference", level=3, color=MID)
for i, c in enumerate(CONTENT["comparisons"], start=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{i}. ")
    r.bold = True
    r = p.add_run(c["q"])
    r.bold = True
    p = doc.add_paragraph(c["a"])
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(6)

doc.save(OUT)
print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")
