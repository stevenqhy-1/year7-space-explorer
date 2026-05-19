"""
Append a 'Do later — Stellarium Web' section with scaffolded tables to a docx.

Usage:
    python3 append-stellarium.py <input.docx> [output.docx]

If output is omitted, the file is updated in place (backup saved alongside).
"""
import sys
from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1a, 0x3a, 0x6a)
MID  = RGBColor(0x2a, 0x4a, 0x7a)
GREY = RGBColor(0x55, 0x55, 0x55)
LINK = RGBColor(0x10, 0x55, 0xb0)

# --- task definitions ---
# Each task: number, title, body sentence, table (header row + N empty rows)
TASKS = [
    {
        "title": "1. The Moon",
        "body": "Find the Moon by clicking it directly or using the search bar at the top of the screen. Record the following values shown in the side panel.",
        "rows": 1,
        "cols": ["Property", "Value"],
        "props": ["Magnitude (brightness)", "Distance", "Radius", "Ra / Dec (location)"],
    },
    {
        "title": "2. Two planets in the Solar System",
        "body": "Pick any two planets (e.g. Venus, Mars, Jupiter). For each, record the values below.",
        "rows": 2,
        "cols": ["Planet", "Magnitude", "Distance (AU)", "Radius", "Ra / Dec"],
        "props": None,
    },
    {
        "title": "3. Two stars",
        "body": "Pick any two stars — Sirius is a good one to start with. Search for them by name and record the values shown.",
        "rows": 2,
        "cols": ["Star", "Magnitude", "Distance", "Ra / Dec"],
        "props": None,
    },
    {
        "title": "4. One galaxy",
        "body": "Search for a galaxy (try 'Andromeda' or 'M31'). Record what Stellarium shows.",
        "rows": 1,
        "cols": ["Galaxy", "Magnitude", "Ra / Dec"],
        "props": None,
    },
    {
        "title": "5. One artificial satellite",
        "body": "Find any artificial satellite — the International Space Station (ISS) is the easiest. Make sure satellites are turned on in the Stellarium settings.",
        "rows": 1,
        "cols": ["Satellite", "Magnitude", "Distance"],
        "props": None,
    },
]


def set_cell_shading(cell, hex_color):
    """Add a background fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="BBBBBB", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tc_pr.append(tcBorders)


def make_hyperlink(paragraph, url, text):
    """Add a real hyperlink run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '1055B0')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_horizontal_rule(doc, color="2A4A7A"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), color)
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)


def append_stellarium_section(doc):
    # Page break to start the new section cleanly
    doc.add_page_break()

    # Section heading
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Do later — Stellarium Web")
    r.bold = True
    r.font.color.rgb = NAVY
    r.font.size = Pt(18)

    # Subtitle / intro
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.add_run("Go to ")
    make_hyperlink(p, "https://stellarium-web.org/", "https://stellarium-web.org/")
    p.add_run(" and have a play around. Then complete the following tasks. Record the values in the tables below.")

    add_horizontal_rule(doc)

    # Each task
    for task in TASKS:
        # Task title
        ph = doc.add_paragraph()
        ph.paragraph_format.space_before = Pt(10)
        ph.paragraph_format.space_after = Pt(2)
        r = ph.add_run(task["title"])
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = NAVY

        # Body sentence
        pb = doc.add_paragraph()
        pb.paragraph_format.space_after = Pt(6)
        r = pb.add_run(task["body"])
        r.font.size = Pt(10.5)
        r.font.color.rgb = GREY

        # Table
        cols = task["cols"]
        # If "props" is given, the table is property/value style (one row per property)
        if task["props"]:
            row_count = len(task["props"]) + 1  # header + rows
        else:
            row_count = task["rows"] + 1  # header + N empty rows

        table = doc.add_table(rows=row_count, cols=len(cols))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = 'Table Grid'

        # Header row
        for i, h in enumerate(cols):
            cell = table.rows[0].cells[i]
            cell.text = ""
            para = cell.paragraphs[0]
            r = para.add_run(h)
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = NAVY
            set_cell_shading(cell, "EEF4FB")

        # Body rows
        if task["props"]:
            for i, prop in enumerate(task["props"]):
                row = table.rows[i + 1]
                # First col = property name
                c0 = row.cells[0]
                c0.text = ""
                p0 = c0.paragraphs[0]
                r = p0.add_run(prop)
                r.font.size = Pt(10.5)
                set_cell_shading(c0, "FAFBFD")
                # Other cols stay blank for student to write
                for ci in range(1, len(cols)):
                    cell = row.cells[ci]
                    cell.text = ""
                    # Give the empty cells a minimum visible height
                    cell.paragraphs[0].add_run(" ")
        else:
            for ri in range(1, row_count):
                row = table.rows[ri]
                for ci in range(len(cols)):
                    cell = row.cells[ci]
                    cell.text = ""
                    cell.paragraphs[0].add_run(" ")

        # Set table column widths so it looks balanced
        # Total ~16 cm
        if task["props"]:
            # property-name col wider
            widths_cm = [4.5] + [(16 - 4.5) / (len(cols) - 1)] * (len(cols) - 1)
        else:
            widths_cm = [16 / len(cols)] * len(cols)
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = Cm(widths_cm[ci])

    # Tip box at end
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Tip:  ")
    r.bold = True
    r.font.color.rgb = NAVY
    p.add_run("Click the search icon (🔍) at the bottom-left of Stellarium Web to look up an object by name. The object's details appear in the side panel after you click it.")


def main():
    if len(sys.argv) < 2:
        print(f"Usage: {sys.argv[0]} <input.docx> [output.docx]")
        sys.exit(1)
    inp = Path(sys.argv[1])
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else inp
    if out == inp:
        backup = inp.with_suffix(inp.suffix + ".bak")
        shutil.copy2(inp, backup)
        print(f"Backup saved: {backup}")

    doc = Document(inp)
    append_stellarium_section(doc)
    doc.save(out)
    print(f"Wrote: {out} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
